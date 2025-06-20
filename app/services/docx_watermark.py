from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_proper_watermark(document, text):
    """
    Add a proper watermark to the Word document on every page
    
    Args:
        document: The python-docx document object
        text: The text for the watermark
        
    Returns:
        document: The modified document object
    """
    # For each section in the document
    for section in document.sections:
        # Get the header
        header = section.header
        
        # Clear existing content in header
        for paragraph in list(header.paragraphs):
            if len(paragraph.text.strip()) > 0:
                p = paragraph._p
                if p.getparent() is not None:
                    p.getparent().remove(p)
        
        # Add a paragraph for the watermark
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a run with the watermark text
        run = paragraph.add_run(text.upper())
        font = run.font
        font.size = Pt(72)  # Large size
        font.bold = True
        font.color.rgb = RGBColor(200, 200, 200)  # Light gray
        
        # Position in the center and rotate
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set the paragraph properties for rotation and positioning
        r = run._r
        rPr = r.get_or_add_rPr()
        
        # Create the vanish effect (text appears in layout but not in printed doc)
        vanish = OxmlElement('w:vanish')
        rPr.append(vanish)
        
        # Create the position effect (centers watermark)
        position = OxmlElement('w:position')
        position.set(qn('w:val'), '-40')  # Negative value moves text up
        rPr.append(position)
        
        # Create watermark effect (positioned in the background)
        webHidden = OxmlElement('w:webHidden')
        rPr.append(webHidden)
        
    return document