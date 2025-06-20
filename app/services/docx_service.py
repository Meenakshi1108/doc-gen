from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import os
from app.services.docx_watermark import add_proper_watermark

def generate_docx(content_html, header_html, footer_html, output_path, watermark=None):
    """
    Generate a DOCX document from HTML content with proper watermark and footer on last page only
    
    Args:
        content_html (str): HTML content for the body
        header_html (str): HTML content for the header
        footer_html (str): HTML content for the footer (last page only)
        output_path (str): Path to save the generated DOCX
        watermark (str, optional): HTML content for watermark
        
    Returns:
        str: Path to the generated document
    """
    # Create a new Document
    doc = Document()
    
    # Set document properties for better formatting
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Parse the content HTML
    soup = BeautifulSoup(content_html, 'html.parser')
    
    # Process each HTML element and add to document
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table']):
        if element.name == 'h1':
            heading = doc.add_heading(element.get_text().strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            heading = doc.add_heading(element.get_text().strip(), level=2)
        elif element.name == 'h3':
            heading = doc.add_heading(element.get_text().strip(), level=3)
        elif element.name == 'p':
            # Check for page break
            if 'page-break-before: always' in element.get('style', ''):
                doc.add_page_break()
            else:
                para = doc.add_paragraph(element.get_text().strip())
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                para = doc.add_paragraph(li.get_text().strip())
                para.style = 'List Bullet' if element.name == 'ul' else 'List Number'
        elif element.name == 'table':
            # Create a table with the right number of rows and columns
            rows = element.find_all('tr')
            if rows:
                cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                
                # Fill the table with data
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < cols:  # Safety check
                            table.cell(i, j).text = cell.get_text().strip()
    
    # Ensure we have a last section for the footer
    # Force creation of a new section for the last page
    doc.add_section(WD_SECTION.NEW_PAGE)
    # Add a small paragraph to the last page to ensure it exists
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header if provided
    if header_html:
        header_soup = BeautifulSoup(header_html, 'html.parser')
        header_text = header_soup.get_text().strip()
        
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Apply formatting
            for run in header_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(119, 119, 119)  # #777
    
    # Add watermark to all pages
    if watermark:
        watermark_text = watermark
        if watermark.startswith('<'):
            watermark_soup = BeautifulSoup(watermark, 'html.parser')
            watermark_text = watermark_soup.get_text().strip()
        
        # Use the proper watermark implementation
        add_proper_watermark(doc, watermark_text)
    
    # Disconnect all footers
    for i, section in enumerate(doc.sections):
        section.footer_distance = Inches(0.5)  # Set consistent footer distance
        section.footer.is_linked_to_previous = False  # Disconnect from previous
        
        # Clear any existing footer content
        for p in list(section.footer.paragraphs):
            p._element.getparent().remove(p._element)
            p._p = None
            p._element = None
    
    # Add footer to last page only
    if footer_html and len(doc.sections) > 0:
        footer_soup = BeautifulSoup(footer_html, 'html.parser')
        footer_text = footer_soup.get_text().strip()
        
        # Get the last section
        last_section = doc.sections[-1]
        
        # Create new footer paragraph
        footer_para = last_section.footer.add_paragraph()
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Style the footer
        for run in footer_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(119, 119, 119)  # #777
        
        # Add a top border to the paragraph
        set_paragraph_border(footer_para)
    
    # Save the document with everything complete
    doc.save(output_path)
    return output_path

def set_paragraph_border(paragraph):
    """
    Add a top border to a paragraph
    
    Args:
        paragraph: The paragraph object to add border to
    """
    p = paragraph._p  # Get the paragraph element
    pPr = p.get_or_add_pPr()  # Get or create paragraph properties
    
    # Create top border element
    top_border = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')  # Border width
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'cccccc')  # Light gray
    top_border.append(top)
    
    # Add border to paragraph properties
    pPr.append(top_border)